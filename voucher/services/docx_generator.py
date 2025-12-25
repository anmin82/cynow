"""
DOCX 견적서/전표 생성 서비스

docxtpl(Jinja2 기반)을 사용하여 DOCX 템플릿에서 문서를 생성합니다.
템플릿 파일만 교체하면 Python 코드 수정 없이 양식 변경이 가능합니다.

사용법:
    from voucher.services import DocxGenerator
    
    generator = DocxGenerator('offer_template.docx')
    output_path = generator.generate(context_data)
"""
import os
import uuid
from datetime import date, datetime
from decimal import Decimal
from typing import Dict, List, Any, Optional
from pathlib import Path

from django.conf import settings
from docxtpl import DocxTemplate


class DocxGenerator:
    """
    DOCX 템플릿 기반 문서 생성기
    
    템플릿 변수 규칙:
    - {{ variable }} : 단순 텍스트 치환
    - {% for item in items %} ... {% endfor %} : 반복 (테이블 행)
    - {% if condition %} ... {% endif %} : 조건부 표시
    """
    
    # 기본 경로 설정
    TEMPLATE_DIR = Path(settings.BASE_DIR) / 'docx_templates'
    OUTPUT_DIR = Path(settings.BASE_DIR) / 'media' / 'generated'
    
    # 통화 기호 매핑
    CURRENCY_SYMBOLS = {
        'KRW': '₩',
        'JPY': '¥',
        'USD': '$',
        'CNY': '¥',
    }
    
    def __init__(self, template_name: str):
        """
        Args:
            template_name: 템플릿 파일명 (예: 'offer_template.docx')
        """
        self.template_path = self.TEMPLATE_DIR / template_name
        
        if not self.template_path.exists():
            raise FileNotFoundError(
                f"템플릿 파일을 찾을 수 없습니다: {self.template_path}"
            )
        
        # 출력 디렉토리 확인
        self.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    
    def generate(
        self, 
        context: Dict[str, Any], 
        output_filename: Optional[str] = None
    ) -> Path:
        """
        템플릿에 데이터를 적용하여 DOCX 파일 생성
        
        Args:
            context: 템플릿 변수 데이터 딕셔너리
            output_filename: 출력 파일명 (기본: UUID 기반 자동 생성)
            
        Returns:
            생성된 파일의 Path 객체
        """
        # 템플릿 로드
        doc = DocxTemplate(self.template_path)
        
        # 컨텍스트 전처리
        processed_context = self._preprocess_context(context)
        
        # 템플릿 렌더링
        doc.render(processed_context)
        
        # 출력 파일명 생성
        if output_filename is None:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"quote_{timestamp}_{uuid.uuid4().hex[:8]}.docx"
        
        output_path = self.OUTPUT_DIR / output_filename
        
        # 파일 저장
        doc.save(output_path)
        
        return output_path
    
    def _preprocess_context(self, context: Dict[str, Any]) -> Dict[str, Any]:
        """
        컨텍스트 데이터 전처리
        - Decimal -> float 변환 (docxtpl 호환)
        - date/datetime -> 문자열 변환
        - 통화 기호 추가
        """
        processed = {}
        
        for key, value in context.items():
            if isinstance(value, Decimal):
                processed[key] = float(value)
            elif isinstance(value, date):
                processed[key] = value.strftime('%Y-%m-%d')
            elif isinstance(value, datetime):
                processed[key] = value.strftime('%Y-%m-%d %H:%M')
            elif isinstance(value, list):
                # 리스트 내 항목들도 전처리
                processed[key] = [
                    self._preprocess_item(item) if isinstance(item, dict) else item
                    for item in value
                ]
            elif isinstance(value, dict):
                processed[key] = self._preprocess_context(value)
            else:
                processed[key] = value
        
        return processed
    
    def _preprocess_item(self, item: Dict[str, Any]) -> Dict[str, Any]:
        """개별 항목 전처리 (테이블 행 데이터)"""
        processed = {}
        
        for key, value in item.items():
            if isinstance(value, Decimal):
                processed[key] = float(value)
            elif isinstance(value, date):
                processed[key] = value.strftime('%Y-%m-%d')
            else:
                processed[key] = value
        
        return processed
    
    def get_currency_symbol(self, currency_code: str) -> str:
        """통화 코드를 기호로 변환"""
        return self.CURRENCY_SYMBOLS.get(currency_code, currency_code)
    
    @classmethod
    def format_number(cls, value: Any, decimal_places: int = 2) -> str:
        """숫자 포맷팅 (천단위 콤마)"""
        if value is None:
            return ''
        try:
            num = float(value)
            if decimal_places == 0:
                return f"{int(num):,}"
            return f"{num:,.{decimal_places}f}"
        except (ValueError, TypeError):
            return str(value)


class QuoteDocxGenerator(DocxGenerator):
    """
    견적서 전용 DOCX 생성기
    
    템플릿 변수 (KDKK 견적서 양식 기준):
    =====================================
    [공급처 정보]
    - supplier_company: 공급처명
    - supplier_address: 주소
    - supplier_ceo: 대표자
    - supplier_tel: TEL
    - supplier_fax: FAX
    - supplier_manager: 담당자명
    - supplier_manager_phone: 담당자 연락처
    - supplier_manager_email: 담당자 이메일
    
    [수신처 정보]
    - receiver_company: 수신처명
    - receiver_address: 주소
    - receiver_ceo: 대표자
    - receiver_tel1: 연락처1
    - receiver_tel2: 연락처2
    - receiver_manager: 담당자명
    - receiver_manager_phone: 담당자 연락처
    - receiver_manager_email: 담당자 이메일
    - gas: 가스 (메모용)
    
    [견적 정보]
    - quote_date: 견적일자
    - quote_no: 견적번호
    - quote_title: 견적건명
    
    [하단 공통 문구]
    - apply_period: 적용기간
    - payment_terms: 거래조건
    - doc_date: 문서 작성일자
    
    [상세 품목 테이블 - items 리스트]
    {%tr for item in items %}
    - item.no: 순번
    - item.gas_name: 가스명
    - item.product_name: 품명
    - item.material_code: 자재코드
    - item.end_user: End User
    - item.packing: 포장
    - item.filling_weight: 충전중량
    - item.currency: 통화
    - item.price_per_kg: 단가(1kg)
    - item.packing_price: 포장단가
    {%tr endfor %}
    =====================================
    """
    
    def __init__(self, template_name: str = 'offer_template.docx'):
        super().__init__(template_name)
    
    def generate_quote(
        self,
        quote_info: Dict[str, Any],
        supplier_info: Dict[str, Any],
        customer_info: Dict[str, Any],
        items: List[Dict[str, Any]],
        footer_info: Optional[Dict[str, Any]] = None,
        output_filename: Optional[str] = None
    ) -> Path:
        """
        견적서 생성
        
        Args:
            quote_info: 견적 기본 정보
            supplier_info: 공급처 정보
            customer_info: 수신처 정보
            items: 품목 리스트
            footer_info: 하단 공통 문구
            output_filename: 출력 파일명
            
        Returns:
            생성된 파일 경로
        """
        # 품목 데이터 가공
        processed_items = []
        for idx, item in enumerate(items, 1):
            processed_item = {
                'no': idx,
                'category': item.get('category', ''),
                'gas_name': item.get('gas_name', ''),
                'product_name': item.get('product_name', ''),
                'material_code': item.get('material_code', ''),
                'end_user': item.get('end_user', ''),
                'packing': item.get('packing', ''),
                'filling_weight': item.get('filling_weight', ''),
                'currency': item.get('currency', 'KRW'),
                'currency_symbol': self.get_currency_symbol(
                    item.get('currency', 'KRW')
                ),
                'price_per_kg': self.format_number(
                    item.get('price_per_kg'), 
                    decimal_places=2
                ),
                'packing_price': self.format_number(
                    item.get('packing_price'),
                    decimal_places=0
                ),
            }
            processed_items.append(processed_item)
        
        # 전체 컨텍스트 구성 (KDKK 견적서 템플릿 변수명 기준)
        context = {
            # 공급처 정보
            'supplier_company': supplier_info.get('name', ''),
            'supplier_address': supplier_info.get('address', ''),
            'supplier_ceo': supplier_info.get('ceo', ''),
            'supplier_tel': supplier_info.get('tel', ''),
            'supplier_fax': supplier_info.get('fax', ''),
            'supplier_manager': supplier_info.get('manager', ''),
            'supplier_manager_phone': supplier_info.get('manager_tel', ''),
            'supplier_manager_email': supplier_info.get('manager_email', ''),
            'no': '',  # 메모용
            
            # 수신처 정보
            'receiver_company': customer_info.get('name', ''),
            'receiver_address': customer_info.get('address', ''),
            'receiver_ceo': customer_info.get('ceo', ''),
            'receiver_tel1': customer_info.get('tel', ''),
            'receiver_tel2': customer_info.get('tel2', ''),
            'receiver_manager': customer_info.get('manager', ''),
            'receiver_manager_phone': customer_info.get('manager_tel', ''),
            'receiver_manager_email': customer_info.get('manager_email', ''),
            'gas': '',  # 메모용
            
            # 견적 정보
            'quote_date': quote_info.get('date', date.today()),
            'quote_no': quote_info.get('no', ''),
            'quote_title': quote_info.get('title', ''),
            
            # 품목 리스트
            'items': processed_items,
            'item_count': len(processed_items),
            
            # 하단 정보
            'apply_period': (footer_info or {}).get('valid_period', ''),
            'payment_terms': (footer_info or {}).get('trade_terms', ''),
            'doc_date': (footer_info or {}).get(
                'document_date', 
                date.today()
            ),
        }
        
        return self.generate(context, output_filename)


def build_quote_context_from_db(quote_id: int) -> Dict[str, Any]:
    """
    DB에서 견적서 데이터를 조회하여 템플릿 컨텍스트 구성
    
    Args:
        quote_id: 견적서 ID
        
    Returns:
        템플릿 컨텍스트 딕셔너리
    """
    from voucher.models import Quote, QuoteItem
    from products.models import ProductCode, ProductPriceHistory
    
    quote = Quote.objects.select_related(
        'customer'
    ).prefetch_related(
        'items__product_code'
    ).get(pk=quote_id)
    
    # 품목 데이터 구성
    items = []
    for item in quote.items.all():
        pc = item.product_code
        items.append({
            'category': item.category or '',
            'gas_name': pc.gas_name if pc else '',
            'product_name': pc.display_name if pc else '',
            'material_code': pc.trade_condition_no if pc else '',
            'end_user': pc.customer_user_name if pc else '',
            'packing': item.packing or '',
            'filling_weight': f'{pc.filling_weight}kg' if pc and pc.filling_weight else '',
            'currency': item.currency,
            'price_per_kg': item.price_per_kg,
            'packing_price': item.packing_price,
        })
    
    return {
        'quote_info': {
            'date': quote.quote_date,
            'no': quote.quote_no,
            'title': quote.title,
        },
        'supplier_info': {
            'name': quote.supplier_name or 'CYNOW Co., Ltd.',
            'address': quote.supplier_address or '',
            'ceo': quote.supplier_ceo or '',
            'tel': quote.supplier_tel or '',
            'fax': quote.supplier_fax or '',
            'manager': quote.supplier_manager or '',
            'manager_tel': quote.supplier_manager_tel or '',
            'manager_email': quote.supplier_manager_email or '',
        },
        'customer_info': {
            'name': quote.customer.name if quote.customer else '',
            'address': quote.customer_address or '',
            'ceo': quote.customer_ceo or '',
            'tel': quote.customer_tel or '',
            'manager': quote.customer_manager or '',
            'manager_tel': quote.customer_manager_tel or '',
            'manager_email': quote.customer_manager_email or '',
        },
        'items': items,
        'footer_info': {
            'valid_period': quote.valid_period or '',
            'trade_terms': quote.trade_terms or '',
            'bank_account': quote.bank_account or '',
            'document_date': quote.created_at.date() if quote.created_at else date.today(),
        },
    }


def generate_price_list_from_products(
    year: int, 
    output_filename: Optional[str] = None
) -> Path:
    """
    제품코드 DB에서 단가표 DOCX 생성 (템플릿 없이 직접 생성)
    
    Args:
        year: 대상 년도
        output_filename: 출력 파일명
        
    Returns:
        생성된 파일 경로
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    
    from products.models import ProductPriceHistory
    
    # 단가 데이터 조회
    prices = ProductPriceHistory.objects.filter(
        effective_date__year=year
    ).select_related('product_code').order_by('product_code__trade_condition_no')
    
    if not prices.exists():
        raise ValueError(f"{year}년 단가 데이터가 없습니다.")
    
    # 문서 생성
    doc = Document()
    
    # 제목
    title = doc.add_heading(f'{year}년도 제품 단가표', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_paragraph(f'KDKK 납품 단가 (적용일: {year}-01-01)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 통화별 그룹핑
    currency_groups = {}
    for price in prices:
        currency = price.currency
        if currency not in currency_groups:
            currency_groups[currency] = []
        currency_groups[currency].append(price)
    
    CURRENCY_SYMBOLS = {'KRW': '₩', 'JPY': '¥', 'USD': '$', 'CNY': '¥'}
    
    # 셀 배경색 설정 함수
    def set_cell_shading(cell, color):
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    # 통화별로 테이블 생성
    for currency, price_list in currency_groups.items():
        symbol = CURRENCY_SYMBOLS.get(currency, '')
        
        doc.add_heading(f'{currency} ({symbol}) 단가', level=1)
        
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 헤더
        headers = ['No', '제품코드', '가스명', '충전량', f'단가 ({symbol}/kg)', '적용일', '비고']
        for i, header in enumerate(table.rows[0].cells):
            header.text = headers[i]
            header.paragraphs[0].runs[0].bold = True
            header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(header, 'D9E2F3')
        
        # 데이터 행
        for idx, price in enumerate(price_list, 1):
            pc = price.product_code
            row = table.add_row()
            cells = row.cells
            
            cells[0].text = str(idx)
            cells[1].text = pc.trade_condition_no
            cells[2].text = pc.gas_name or ''
            cells[3].text = f'{pc.filling_weight or 0}kg'
            cells[4].text = f'{price.price_per_kg:,.2f}'
            cells[5].text = str(price.effective_date)
            cells[6].text = price.note or ''
            
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
    
    # 푸터
    doc.add_paragraph()
    footer = doc.add_paragraph(f'발행일: {date.today().strftime("%Y-%m-%d")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    company = doc.add_paragraph('CYNOW Co., Ltd.')
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 저장
    output_dir = Path(settings.BASE_DIR) / 'media' / 'generated'
    output_dir.mkdir(parents=True, exist_ok=True)
    
    if output_filename is None:
        output_filename = f'{year}년단가표_KDKK.docx'
    
    output_path = output_dir / output_filename
    doc.save(output_path)
    
    return output_path

