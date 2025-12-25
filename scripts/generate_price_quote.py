#!/usr/bin/env python
"""
2026년 KDKK 단가표 DOCX 생성 스크립트

사용법:
  python scripts/generate_price_quote.py [년도]
  예: python scripts/generate_price_quote.py 2026
"""
import os
import sys
import django
from datetime import date
from decimal import Decimal

# Django 설정
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')
django.setup()

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from products.models import ProductPriceHistory, ProductCode

CURRENCY_SYMBOLS = {
    'KRW': '₩',
    'JPY': '¥',
    'USD': '$',
    'CNY': '¥',
}

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def generate_price_quote(year: int):
    """단가표 DOCX 생성"""
    
    # 해당 년도 단가 조회
    prices = ProductPriceHistory.objects.filter(
        effective_date__year=year
    ).select_related('product_code').order_by('product_code__trade_condition_no')
    
    if not prices.exists():
        print(f"[ERROR] {year}년 단가 데이터가 없습니다.")
        return None
    
    # 문서 생성
    doc = Document()
    
    # 제목
    title = doc.add_heading(f'{year}년도 제품 단가표', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 부제목
    subtitle = doc.add_paragraph(f'KDKK 납품 단가 (적용일: {year}-01-01)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # 빈 줄
    
    # 통화별 그룹핑
    currency_groups = {}
    for price in prices:
        currency = price.currency
        if currency not in currency_groups:
            currency_groups[currency] = []
        currency_groups[currency].append(price)
    
    # 통화별로 테이블 생성
    for currency, price_list in currency_groups.items():
        symbol = CURRENCY_SYMBOLS.get(currency, '')
        
        # 통화 섹션 헤딩
        doc.add_heading(f'{currency} ({symbol}) 단가', level=1)
        
        # 테이블 생성 (7열)
        table = doc.add_table(rows=1, cols=7)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 헤더 행
        header_cells = table.rows[0].cells
        headers = ['No', '제품코드', '가스명', '충전량', f'단가 ({symbol}/kg)', '적용일', '비고']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_cell_shading(header_cells[i], 'D9E2F3')  # 연한 파란색
        
        # 데이터 행
        for idx, price in enumerate(price_list, 1):
            pc = price.product_code
            row = table.add_row()
            cells = row.cells
            
            cells[0].text = str(idx)
            cells[1].text = pc.trade_condition_no
            cells[2].text = pc.gas_name or ''
            cells[3].text = f'{pc.filling_weight or 0}kg'
            cells[4].text = f'{price.price_per_kg:,.2f}'
            cells[5].text = str(price.effective_date)
            cells[6].text = price.note or ''
            
            # 정렬
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # 테이블 사이 간격
    
    # 푸터
    doc.add_paragraph()
    footer = doc.add_paragraph(f'발행일: {date.today().strftime("%Y-%m-%d")}')
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    company = doc.add_paragraph('CYNOW Co., Ltd.')
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 파일 저장
    output_path = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        f'{year}년단가표_KDKK.docx'
    )
    doc.save(output_path)
    
    print(f'[OK] Price quote generated: {output_path}')
    print(f'     Total: {prices.count()} items')
    print(f'     By currency: {", ".join([f"{k}: {len(v)}" for k, v in currency_groups.items()])}')
    
    return output_path


if __name__ == '__main__':
    year = int(sys.argv[1]) if len(sys.argv) > 1 else 2026
    generate_price_quote(year)

