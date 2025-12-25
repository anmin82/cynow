#!/usr/bin/env python3
"""
견적서 템플릿 생성 스크립트

기존 견적서 양식을 기반으로 Jinja2 변수가 적용된 템플릿을 생성합니다.
"""
import os
import sys
from pathlib import Path

# Django 설정
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'config.settings')

import django
django.setup()

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """셀 테두리 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            element.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            element.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def create_offer_template():
    """견적서 템플릿 생성"""
    doc = Document()
    
    # 페이지 설정
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # 제목
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('견 적 서')
    run.bold = True
    run.font.size = Pt(24)
    run.font.name = '맑은 고딕'
    
    doc.add_paragraph()  # 빈 줄
    
    # ========================================
    # 상단 정보 테이블 (공급자 + 수신처)
    # ========================================
    info_table = doc.add_table(rows=8, cols=4)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 공급자 정보 (왼쪽)
    info_table.cell(0, 0).text = '공급자'
    info_table.cell(0, 1).text = '{{supplier_company}}'
    info_table.cell(1, 0).text = '주소'
    info_table.cell(1, 1).text = '{{supplier_address}}'
    info_table.cell(2, 0).text = '대표이사'
    info_table.cell(2, 1).text = '{{supplier_ceo}}'
    info_table.cell(3, 0).text = 'TEL'
    info_table.cell(3, 1).text = '{{supplier_tel}}'
    info_table.cell(4, 0).text = 'FAX'
    info_table.cell(4, 1).text = '{{supplier_fax}}'
    info_table.cell(5, 0).text = '담당자'
    info_table.cell(5, 1).text = '{{supplier_manager}}'
    info_table.cell(6, 0).text = '연락처'
    info_table.cell(6, 1).text = '{{supplier_manager_phone}}'
    info_table.cell(7, 0).text = 'E-mail'
    info_table.cell(7, 1).text = '{{supplier_manager_email}}'
    
    # 수신처 정보 (오른쪽)
    info_table.cell(0, 2).text = '수신'
    info_table.cell(0, 3).text = '{{receiver_company}}'
    info_table.cell(1, 2).text = '주소'
    info_table.cell(1, 3).text = '{{receiver_address}}'
    info_table.cell(2, 2).text = '대표이사'
    info_table.cell(2, 3).text = '{{receiver_ceo}}'
    info_table.cell(3, 2).text = 'TEL'
    info_table.cell(3, 3).text = '{{receiver_tel}}'
    info_table.cell(4, 2).text = '담당자'
    info_table.cell(4, 3).text = '{{receiver_manager}}'
    info_table.cell(5, 2).text = '연락처'
    info_table.cell(5, 3).text = '{{receiver_manager_phone}}'
    info_table.cell(6, 2).text = 'E-mail'
    info_table.cell(6, 3).text = '{{receiver_manager_email}}'
    info_table.cell(7, 2).text = ''
    info_table.cell(7, 3).text = ''
    
    doc.add_paragraph()  # 빈 줄
    
    # ========================================
    # 견적 정보
    # ========================================
    quote_info = doc.add_paragraph()
    quote_info.add_run('견적번호: ').bold = True
    quote_info.add_run('{{quote_no}}')
    quote_info.add_run('    ')
    quote_info.add_run('견적일자: ').bold = True
    quote_info.add_run('{{quote_date}}')
    quote_info.add_run('    ')
    quote_info.add_run('건명: ').bold = True
    quote_info.add_run('{{quote_title}}')
    
    doc.add_paragraph()  # 빈 줄
    
    # ========================================
    # 품목 테이블
    # ========================================
    item_table = doc.add_table(rows=2, cols=10)
    item_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 헤더
    headers = ['No', '가스명', '제품명', '자재코드', 'End User', 
               '용기', '충전량(kg)', '통화', '단가(/kg)', '용기단가']
    for i, header in enumerate(headers):
        cell = item_table.cell(0, i)
        cell.text = header
        # 헤더 스타일
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
    
    # 데이터 행 (반복 템플릿)
    # 첫 번째 셀에 for 시작, 마지막 셀에 endfor
    item_table.cell(1, 0).text = '{%tr for item in items %}{{item.no}}'
    item_table.cell(1, 1).text = '{{item.gas_name}}'
    item_table.cell(1, 2).text = '{{item.product_name}}'
    item_table.cell(1, 3).text = '{{item.material_code}}'
    item_table.cell(1, 4).text = '{{item.end_user}}'
    item_table.cell(1, 5).text = '{{item.packing}}'
    item_table.cell(1, 6).text = '{{item.filling_weight}}'
    item_table.cell(1, 7).text = '{{item.currency}}'
    item_table.cell(1, 8).text = '{{item.price_per_kg}}'
    item_table.cell(1, 9).text = '{{item.packing_price}}{%tr endfor %}'
    
    doc.add_paragraph()  # 빈 줄
    
    # ========================================
    # 하단 정보
    # ========================================
    footer_table = doc.add_table(rows=3, cols=2)
    footer_table.cell(0, 0).text = '적용기간'
    footer_table.cell(0, 1).text = '{{apply_period}}'
    footer_table.cell(1, 0).text = '결제조건'
    footer_table.cell(1, 1).text = '{{payment_terms}}'
    footer_table.cell(2, 0).text = '작성일자'
    footer_table.cell(2, 1).text = '{{doc_date}}'
    
    doc.add_paragraph()  # 빈 줄
    
    # 서명란
    sign = doc.add_paragraph()
    sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign.add_run('{{doc_date}}')
    
    sign2 = doc.add_paragraph()
    sign2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign2.add_run('{{supplier_company}}')
    
    sign3 = doc.add_paragraph()
    sign3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign3.add_run('대표이사 {{supplier_ceo}} (인)')
    
    # 저장
    output_path = Path(__file__).resolve().parent.parent / 'docx_templates' / 'offer_template.docx'
    doc.save(output_path)
    print(f'✅ 템플릿 생성 완료: {output_path}')
    
    return output_path


if __name__ == '__main__':
    create_offer_template()

